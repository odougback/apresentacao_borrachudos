"""Gera relatorio.docx a partir do conteúdo do relatório técnico de vazão.

Uso:
    python3 gerar_docx.py

Saída: vazao/relatorio.docx (na mesma pasta deste script).
"""

from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


HERE = Path(__file__).resolve().parent

# Paleta visual do relatório web
RED = RGBColor(0xC8, 0x39, 0x2F)
RED_DARK = RGBColor(0xA6, 0x2B, 0x22)
INK = RGBColor(0x1A, 0x14, 0x14)
MUTED = RGBColor(0x6B, 0x5E, 0x5B)
GREEN = RGBColor(0x10, 0xB9, 0x81)
ORANGE = RGBColor(0xD9, 0x77, 0x06)


def shade(cell_or_paragraph, color_hex: str) -> None:
    """Aplica preenchimento (shading) em célula de tabela."""
    tc_pr = cell_or_paragraph._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int, color: RGBColor = INK,
                size: Optional[int] = None) -> None:
    """Cabeçalho com cor/tamanho personalizados."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.bold = True
    run.font.color.rgb = color
    sizes = {1: 22, 2: 18, 3: 14, 4: 12}
    run.font.size = Pt(size or sizes.get(level, 12))
    # marca como heading para o sumário do Word
    p.style = doc.styles[f"Heading {min(level, 4)}"]
    # sobrescreve cor do estilo
    run.font.color.rgb = color


def add_paragraph(doc: Document, text: str, *, bold: bool = False,
                  italic: bool = False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  size: int = 11, color: RGBColor = INK) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_rich_paragraph(doc: Document, fragments: list[tuple[str, dict]],
                       align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    """Parágrafo com runs múltiplos. Cada fragmento: (texto, {bold, italic, color, size})."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    for text, opts in fragments:
        run = p.add_run(text)
        run.font.size = Pt(opts.get("size", 11))
        run.font.bold = opts.get("bold", False)
        run.font.italic = opts.get("italic", False)
        run.font.color.rgb = opts.get("color", INK)


def add_bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.runs[0] if p.runs else p.add_run()
        run.text = item
        run.font.size = Pt(11)
        run.font.color.rgb = INK


def add_formula_box(doc: Document, equation: str, variables: list[str]) -> None:
    """Caixa escura com fórmula centralizada e lista de variáveis."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    shade(cell, "1A1414")

    # Equação centralizada
    eq_p = cell.paragraphs[0]
    eq_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq_p.paragraph_format.space_before = Pt(8)
    eq_p.paragraph_format.space_after = Pt(8)
    eq_run = eq_p.add_run(equation)
    eq_run.font.size = Pt(16)
    eq_run.font.bold = True
    eq_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Linha separadora
    sep_p = cell.add_paragraph()
    sep_p.paragraph_format.space_before = Pt(0)
    sep_p.paragraph_format.space_after = Pt(4)
    sep_run = sep_p.add_run("─" * 50)
    sep_run.font.size = Pt(8)
    sep_run.font.color.rgb = RGBColor(0x55, 0x44, 0x44)

    # Variáveis
    for var in variables:
        v_p = cell.add_paragraph()
        v_p.paragraph_format.space_after = Pt(2)
        v_p.paragraph_format.left_indent = Cm(0.4)
        v_run = v_p.add_run("• " + var)
        v_run.font.size = Pt(10)
        v_run.font.color.rgb = RGBColor(0xE0, 0xE0, 0xE0)

    # spacing pós-tabela
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)


def add_formula_image(doc: Document, path: Path, width_cm: float = 16) -> None:
    """Embute o PNG do formula-card centralizado, sem legenda."""
    if not path.exists():
        add_paragraph(doc, f"[imagem ausente: {path.name}]", italic=True, color=MUTED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))


def add_callout_warning(doc: Document, title: str, paragraphs: list[str]) -> None:
    """Caixa de observação técnica (fundo claro, borda lateral laranja)."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    shade(cell, "FEF7F4")

    # Título
    t_p = cell.paragraphs[0]
    t_p.paragraph_format.space_before = Pt(6)
    t_p.paragraph_format.space_after = Pt(6)
    t_run = t_p.add_run(title.upper())
    t_run.font.bold = True
    t_run.font.size = Pt(10)
    t_run.font.color.rgb = ORANGE

    for text in paragraphs:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.color.rgb = INK

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)


def add_image(doc: Document, path: Path, caption: str, width_cm: float = 16) -> None:
    if not path.exists():
        add_paragraph(doc, f"[imagem ausente: {path.name}]", italic=True, color=MUTED)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Cm(width_cm))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    cap_run = cap.add_run(caption)
    cap_run.font.italic = True
    cap_run.font.size = Pt(10)
    cap_run.font.color.rgb = MUTED


def add_dose_card(doc: Document) -> None:
    """Caixa verde de destaque com 'DOSE CALCULADA ≈ 58 ml'."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    shade(cell, "10B981")

    label = cell.paragraphs[0]
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label.paragraph_format.space_before = Pt(8)
    label.paragraph_format.space_after = Pt(2)
    l_run = label.add_run("DOSE CALCULADA")
    l_run.font.size = Pt(11)
    l_run.font.bold = True
    l_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    value = cell.add_paragraph()
    value.alignment = WD_ALIGN_PARAGRAPH.CENTER
    value.paragraph_format.space_after = Pt(8)
    v_run = value.add_run("≈ 58 ml")
    v_run.font.size = Pt(28)
    v_run.font.bold = True
    v_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def add_parameter_table(doc: Document) -> None:
    rows = [
        ("Largura", "1,288 m", "item 2.2.1"),
        ("Profundidade", "0,1433 m", "item 2.2.2"),
        ("Tempo médio", "40,67 s", "item 2.2.3"),
        ("Concentração", "25 ppm", "padrão"),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Light Grid Accent 1"
    headers = ("Parâmetro", "Valor", "Origem")
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(c, "1A1414")
    for r, (param, val, src) in enumerate(rows, start=1):
        for i, txt in enumerate((param, val, src)):
            c = table.rows[r].cells[i]
            c.text = ""
            run = c.paragraphs[0].add_run(txt)
            run.font.size = Pt(10.5)
            if i == 1:
                run.font.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build_cover(doc: Document) -> None:
    # Tag superior
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag.paragraph_format.space_before = Pt(80)
    tag.paragraph_format.space_after = Pt(20)
    tag_run = tag.add_run("RELATÓRIO TÉCNICO")
    tag_run.font.bold = True
    tag_run.font.size = Pt(11)
    tag_run.font.color.rgb = RED

    # Título
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    t_run = title.add_run("Avaliação de Pontos do Controle Larval")
    t_run.font.bold = True
    t_run.font.size = Pt(28)
    t_run.font.color.rgb = INK

    # Versão
    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver.paragraph_format.space_after = Pt(28)
    v_run = ver.add_run("VERSÃO 05")
    v_run.font.bold = True
    v_run.font.size = Pt(13)
    v_run.font.color.rgb = RED

    # Metadata box
    meta = doc.add_table(rows=2, cols=2)
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.autofit = False
    rows = [("Pregão Eletrônico", "590/2023"),
            ("Termo de Contrato", "650/2024")]
    for r, (label, value) in enumerate(rows):
        lc, vc = meta.rows[r].cells
        lc.text = ""
        vc.text = ""
        l_run = lc.paragraphs[0].add_run(label)
        l_run.font.size = Pt(11)
        l_run.font.color.rgb = MUTED
        v_run = vc.paragraphs[0].add_run(value)
        v_run.font.size = Pt(11)
        v_run.font.bold = True
        v_run.font.color.rgb = INK
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Cliente
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(28)
    cli = doc.add_paragraph()
    cli.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cli.paragraph_format.space_after = Pt(2)
    c_run = cli.add_run("Município de Joinville")
    c_run.font.bold = True
    c_run.font.size = Pt(16)
    c_run.font.color.rgb = INK

    org = doc.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org.paragraph_format.space_after = Pt(28)
    o_run = org.add_run("Secretaria de Meio Ambiente")
    o_run.font.size = Pt(11)
    o_run.font.color.rgb = MUTED

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d_run = date.add_run("ABRIL / 2026")
    d_run.font.bold = True
    d_run.font.size = Pt(11)
    d_run.font.color.rgb = MUTED
    d_run.font.name = "Calibri"

    # Quebra de página
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build_toc(doc: Document) -> None:
    add_heading(doc, "Sumário", level=2, color=RED)
    items = [
        ("1.", "Introdução"),
        ("2.", "Metodologia para Determinação de Pontos"),
        ("2.1", "Levantamento e avaliação de campo"),
        ("2.2", "Medição de Vazão"),
        ("2.2.1", "Determinação da largura média"),
        ("2.2.2", "Determinação da profundidade média"),
        ("2.2.3", "Velocidade superficial"),
        ("2.2.4", "Dimensionamento da dose de larvicida"),
        ("2.3", "Identificação dos Pontos"),
        ("3.", "Execução das Medições de Vazão em Campo"),
        ("3.1", "Considerações iniciais"),
        ("3.2", "Regiões contempladas"),
        ("3.3", "Pontos com medição direta de vazão"),
        ("3.4", "Apresentação dos pontos sem medição de vazão"),
        ("3.5", "Comprovação técnica e atualizações futuras"),
    ]
    for num, label in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        # indenta sub-itens
        depth = num.count(".")
        if num.endswith("."):
            depth -= 1
        p.paragraph_format.left_indent = Cm(0.5 * depth)
        n = p.add_run(f"{num}  ")
        n.font.bold = True
        n.font.color.rgb = RED
        n.font.size = Pt(11)
        t = p.add_run(label)
        t.font.size = Pt(11)
        t.font.color.rgb = INK

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build_chapter_1(doc: Document) -> None:
    add_heading(doc, "1. Introdução", level=1, color=INK)

    paragraphs = [
        "A avaliação dos pontos de controle larval constitui uma etapa fundamental no manejo de simulídeos (borrachudos), uma vez que permite definir, com precisão técnica, os locais estratégicos para aplicação do larvicida biológico, visando a máxima eficiência no controle das formas imaturas (larvas).",
        "Esse processo tem como principal objetivo identificar, ajustar e validar os pontos de intervenção ao longo dos cursos d'água, garantindo que o produto seja adequadamente distribuído e alcance as áreas com presença de larvas, considerando as características hídricas e ambientais de cada trecho.",
        "A definição dos pontos de controle deve levar em consideração, de forma integrada, a presença de insetos adultos na área e a ocorrência de formas imaturas (larvas) nos criadouros, permitindo correlacionar os focos de infestação com os locais de desenvolvimento do vetor.",
        "A avaliação contempla, de forma técnica, os seguintes aspectos:",
    ]
    for p in paragraphs:
        add_paragraph(doc, p)

    add_bullet_list(doc, [
        "Análise e definição dos pontos de aplicação, incluindo a possibilidade de realocação de pontos existentes ou implantação de novos pontos, conforme as condições observadas em campo;",
        "Determinação da vazão dos cursos d'água, parâmetro essencial para o dimensionamento da dose de larvicida a ser aplicada.",
    ])

    add_paragraph(doc,
        "A execução adequada dessas atividades é determinante para a eficiência do tratamento, pois influencia diretamente a dispersão do produto, o alcance da aplicação e a efetividade no controle das larvas. Além disso, contribui para a otimização dos recursos operacionais e para a sustentabilidade das ações de controle, reduzindo desperdícios e minimizando impactos ambientais.")


def build_chapter_2(doc: Document) -> None:
    add_heading(doc, "2. Metodologia para Determinação de Pontos", level=1, color=INK)
    add_paragraph(doc,
        "A metodologia adotada para a avaliação dos pontos de controle larval de simulídeos baseia-se em procedimentos técnicos de campo voltados à identificação, caracterização e dimensionamento dos locais de aplicação do larvicida biológico.")
    add_paragraph(doc,
        "As atividades são realizadas diretamente nos cursos d'água, considerando suas características hídricas, ambientais e biológicas, com o objetivo de garantir maior precisão na definição dos pontos e na eficiência do controle.")
    add_paragraph(doc, "A metodologia compreende as seguintes etapas:")

    # 2.1
    add_heading(doc, "2.1 Levantamento e avaliação de campo", level=2)
    add_paragraph(doc,
        "Consiste na inspeção detalhada dos cursos d'água, com foco na identificação de áreas com presença de larvas e na verificação da ocorrência de insetos adultos nas proximidades.")
    add_paragraph(doc, "Durante essa etapa, são avaliados:")
    add_bullet_list(doc, [
        "Presença e intensidade de infestação larval;",
        "Ocorrência de insetos adultos na área;",
        "Características do curso d'água (largura, profundidade, velocidade aparente da corrente);",
        "Presença de substratos favoráveis ao desenvolvimento larval (rochas, vegetação, galhos e materiais submersos);",
        "Condições hídricas do trecho (turbulência, remansos, obstáculos e acúmulo de matéria orgânica).",
    ])
    add_paragraph(doc,
        "Com base nessas informações, são definidos, ajustados ou implantados novos pontos de aplicação, visando garantir cobertura adequada das áreas infestadas.")

    # 2.2
    add_heading(doc, "2.2 Medição de Vazão", level=2)
    add_paragraph(doc,
        "Quando há necessidade de abertura de novos pontos em áreas ainda não contempladas pelo Programa de Controle de Simulídeos, a determinação da vazão do curso d'água constitui etapa fundamental para o correto dimensionamento da dose de larvicida a ser aplicada.")
    add_paragraph(doc,
        "Sempre que houver acesso seguro e condições adequadas de campo, a vazão de rios, córregos ou canais deve ser estimada por meio da medição direta de parâmetros hídricos básicos, sendo esta a metodologia recomendada para garantir maior precisão nos cálculos e na eficiência da aplicação.")
    add_paragraph(doc,
        "A vazão é determinada a partir da multiplicação de quatro parâmetros: largura média do canal, profundidade média da lâmina d'água, velocidade superficial da corrente e um coeficiente de correção, conforme a seguinte expressão:")

    add_formula_image(doc, HERE / "formula_vazao.png")

    add_paragraph(doc,
        "O coeficiente de correção (C) é aplicado para converter a velocidade superficial — medida pelo método do flutuador — em velocidade média ao longo da seção transversal do curso d'água. Como o atrito com o leito e as margens faz com que a água do fundo e das laterais escoe mais lentamente que a da superfície, a velocidade superficial é sempre maior que a velocidade média efetiva. O coeficiente, geralmente situado entre 0,80 e 0,90, ajusta essa diferença, sendo o valor de 0,85 amplamente adotado em programas de controle de simulídeos para leitos naturais com características hidráulicas regulares.")
    add_paragraph(doc,
        "A aplicação dessa metodologia permite estimar o volume de água que escoa no trecho avaliado em determinado intervalo de tempo, sendo um parâmetro essencial para o cálculo da dosagem adequada do larvicida. Dessa forma, garante-se maior eficiência no tratamento, melhor dispersão do produto e cobertura adequada das áreas com presença de larvas.")
    add_paragraph(doc,
        "Em situações específicas onde a medição direta não possa ser realizada de forma segura ou operacionalmente viável, o dimensionamento da dose poderá ser definido com base em critérios técnicos empíricos. Nesses casos, a estimativa considera a experiência do profissional responsável, associada à avaliação visual das características do curso d'água, tais como largura aparente, velocidade da corrente, presença de matéria orgânica, nível de infestação e comportamento hídrico do trecho (turbulência, remansos e obstáculos).")
    add_paragraph(doc,
        "Ainda que essa abordagem seja aplicada em campo, recomenda-se que seja utilizada de forma excepcional, sendo sempre acompanhada de monitoramento posterior, com o objetivo de validar a eficácia do tratamento e permitir ajustes nas aplicações subsequentes.")

    # 2.2.1
    add_heading(doc, "2.2.1 Determinação da largura média", level=3)
    for txt in [
        "A determinação da largura média do curso d'água deve ser realizada por meio de medições sucessivas ao longo de um trecho representativo, preferencialmente com extensão de 10 metros.",
        "Para isso, recomenda-se medir a largura do córrego em intervalos regulares de 1 metro, totalizando múltiplos pontos de amostragem ao longo do trecho selecionado. Essas medições devem ser realizadas de margem a margem, em seções perpendiculares ao fluxo da água, buscando representar fielmente as variações do canal.",
        "Após a coleta dos dados, os valores obtidos devem ser somados e divididos pelo número total de medições realizadas, resultando na largura média do curso d'água, expressa em metros.",
        "Esse parâmetro é fundamental para o cálculo da vazão, sendo diretamente proporcional ao volume de água transportado no trecho avaliado e, consequentemente, ao dimensionamento da dose de larvicida a ser aplicada.",
    ]:
        add_paragraph(doc, txt)

    add_image(doc, HERE / "svg_largura.png",
              "Figura 1 — Determinação da largura média: 10 medições perpendiculares ao fluxo, espaçadas a cada 1 metro. Larguras expressas em centímetros.")

    add_paragraph(doc, "Exemplo de cálculo:", bold=True)
    add_bullet_list(doc, [
        "Larguras medidas (cm): 183, 170, 125, 113, 135, 127, 130, 115, 100, 90",
        "Soma: 1.288 cm",
        "Média: 1.288 ÷ 10 = 128,8 cm = 1,288 m",
    ])

    # 2.2.2
    add_heading(doc, "2.2.2 Determinação da profundidade média", level=3)
    for txt in [
        "A profundidade média do curso d'água deve ser determinada no mesmo trecho utilizado para a medição da largura, garantindo a consistência dos dados hídricos.",
        "Para cada seção de largura previamente definida (a cada 1 metro ao longo do trecho), devem ser realizadas medições de profundidade em três pontos distintos: próximo à margem esquerda, no centro do canal e próximo à margem direita. Essa distribuição permite representar de forma mais fiel as variações do leito do curso d'água.",
        "As medições devem ser executadas de maneira uniforme ao longo de todas as seções, contemplando tanto regiões mais rasas quanto mais profundas, evitando distorções no resultado final.",
        "Após a coleta dos dados, todos os valores de profundidade obtidos devem ser somados e divididos pelo número total de medições realizadas, resultando na profundidade média do curso d'água, expressa em metros.",
        "Esse parâmetro, associado à largura média e à velocidade da corrente, é fundamental para o cálculo da vazão e, consequentemente, para o correto dimensionamento da dose de larvicida a ser aplicada.",
    ]:
        add_paragraph(doc, txt)

    # Figura 2 — vista superior + corte transversal (dois SVGs)
    add_image(doc, HERE / "svg_profundidade_top.png",
              "Vista superior — 30 pontos (3 por seção × 10 seções).")
    # Divisor entre as duas vistas
    div = doc.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div.paragraph_format.space_before = Pt(4)
    div.paragraph_format.space_after = Pt(4)
    div_run = div.add_run("VISTA EM CORTE (SEÇÃO TRANSVERSAL) — EXEMPLO DA SEÇÃO 4")
    div_run.font.size = Pt(9)
    div_run.font.bold = True
    div_run.font.color.rgb = RED
    add_image(doc, HERE / "svg_profundidade_cross.png",
              "Figura 2 — Determinação da profundidade média: 3 medições por seção (esquerda, centro, direita) ao longo de 10 seções, totalizando 30 pontos. Profundidades expressas em centímetros.")

    add_paragraph(doc, "Exemplo de cálculo:", bold=True)
    add_bullet_list(doc, [
        "Soma das profundidades: Σ Pe (107 cm) + Σ Pc (219 cm) + Σ Pd (104 cm) = 430 cm",
        "Média: 430 ÷ 30 medições = 14,33 cm = 0,1433 m",
    ])

    # 2.2.3
    add_heading(doc, "2.2.3 Velocidade superficial", level=3)
    for txt in [
        "A velocidade superficial do curso d'água deve ser determinada no mesmo trecho onde foram realizadas as medições de largura e profundidade médias, garantindo a representatividade dos dados utilizados no cálculo da vazão.",
        "Para essa medição, utiliza-se um objeto flutuante (como uma esfera parcialmente preenchida, por exemplo, uma bola de ping-pong ou similar), que deve ser lançado na correnteza em um ponto inicial previamente definido.",
        "Em seguida, mede-se o tempo necessário para que o objeto percorra uma distância conhecida, recomendando-se um trecho de 10 metros com fluxo relativamente homogêneo. Esse procedimento deve ser repetido mais de uma vez, a fim de reduzir possíveis interferências e obter maior confiabilidade nos resultados.",
        "A velocidade superficial é então determinada pela relação entre a distância percorrida e o tempo médio obtido nas medições, conforme a expressão:",
    ]:
        add_paragraph(doc, txt)

    add_formula_image(doc, HERE / "formula_velocidade.png")

    add_paragraph(doc,
        "A correta determinação da velocidade é fundamental para o cálculo da vazão, influenciando diretamente o dimensionamento da dose de larvicida e a eficiência da aplicação.")

    add_image(doc, HERE / "svg_velocidade.png",
              "Figura 3 — Determinação da velocidade superficial: cronometragem do tempo que o flutuador percorre uma distância conhecida (10 m). O procedimento é repetido ao menos 3 vezes para obter o tempo médio.")

    add_paragraph(doc, "Exemplo de cálculo:", bold=True)
    add_bullet_list(doc, [
        "Tentativas: T₁ = 39 s · T₂ = 43 s · T₃ = 40 s",
        "Tempo médio: (39 + 43 + 40) ÷ 3 = 40,67 s",
        "Velocidade: Vs = 10 ÷ 40,67 = 0,25 m/s",
    ])

    # 2.2.4
    add_heading(doc, "2.2.4 Dimensionamento da dose de larvicida", level=3)
    add_paragraph(doc,
        "O dimensionamento da dose de VectoBac® 12AS para o controle de simulídeos é realizado com base nos dados obtidos em campo, sendo utilizado, como ferramenta de apoio, planilha técnica fornecida pelo fabricante do produto. Essa planilha automatiza o cálculo da dose a partir dos parâmetros hídricos medidos, garantindo maior padronização e confiabilidade nos resultados.")
    add_paragraph(doc,
        "O cálculo da dose é fundamentado nas medições de largura média, profundidade média, tempo de escoamento da água e no mesmo coeficiente de correção utilizado no cálculo da vazão, seguindo a seguinte expressão:")

    add_formula_image(doc, HERE / "formula_dose.png")

    add_paragraph(doc,
        "O coeficiente de correção (C) é mantido no cálculo da dose pelas mesmas razões descritas no item 2.2: como o termo (600 / TEMPO) é derivado da velocidade superficial do flutuador, sua aplicação direta superestimaria o volume de água efetivamente escoado na seção. A inclusão de C garante a coerência entre o dimensionamento da dose e a vazão real do trecho, evitando subdosagem ou desperdício de produto.")

    add_rich_paragraph(doc, [
        ("A concentração de ", {}),
        ("25 ppm", {"bold": True}),
        (" é adotada como padrão no dimensionamento da dose para a aplicação de VectoBac® 12AS no controle de simulídeos, parâmetro consolidado pela planilha técnica do fabricante. Aplicando a fórmula com os dados medidos nas seções 2.2.1, 2.2.2 e 2.2.3 (L = 1,288 m; P = 0,1433 m; T = 40,67 s), obtém-se a dose detalhada a seguir:", {}),
    ])

    add_image(doc, HERE / "svg_dose.png",
              "Figura 4 — Exemplo do dimensionamento da dose com os dados medidos nas seções 2.2.1, 2.2.2 e 2.2.3. Para um trecho com largura média de 1,288 m, profundidade média de 0,1433 m e tempo médio de 40,67 s, à concentração padrão de 25 ppm, a dose calculada de VectoBac® 12AS é de aproximadamente 58 ml para o trecho avaliado, em conformidade com a planilha técnica do fabricante.")

    add_callout_warning(doc, "Observação técnica", [
        "Em situações nas quais não seja possível realizar a medição direta da vazão do curso d'água — seja por limitações de acesso, segurança ou condições operacionais — o dimensionamento da dose do larvicida deve ser realizado com base em critérios técnicos empíricos.",
        "Nesses casos, a definição da dosagem é conduzida pelo profissional responsável, considerando sua experiência de campo e a avaliação visual das características do trecho, tais como largura aparente do canal, profundidade estimada, velocidade perceptível da corrente, presença de matéria orgânica em suspensão, nível de infestação larval e condições hídricas (turbulência, remansos e obstáculos).",
        "Essa abordagem, embora menos precisa do que a medição direta, é amplamente empregada em programas de controle em áreas de difícil acesso, permitindo a continuidade das ações operacionais.",
        "Ressalta-se que, nessas situações, é indispensável a realização de monitoramento posterior, com o objetivo de avaliar a eficácia do tratamento e promover ajustes nas doses e nos pontos de aplicação em intervenções subsequentes, garantindo a efetividade do controle.",
    ])

    # 2.3
    add_heading(doc, "2.3 Identificação dos Pontos", level=2)
    for txt in [
        "Após a definição dos pontos de aplicação e o dimensionamento da dose de larvicida a ser utilizada em cada local, procede-se à identificação física dos pontos em campo, com o objetivo de garantir sua correta localização, padronização das aplicações e rastreabilidade das intervenções.",
        "A identificação é realizada por meio da instalação de marcos visuais no local, como placas fixas, devidamente posicionadas em áreas de fácil visualização e acesso. Essas identificações devem conter informações essenciais, como código do ponto.",
        "Paralelamente, cada ponto deve estar vinculado a uma ficha de controle, na qual são registradas todas as informações operacionais como: dose aplicada, datas de aplicação.",
        "Esse procedimento permite:",
    ]:
        add_paragraph(doc, txt)
    add_bullet_list(doc, [
        "Padronizar as operações de campo;",
        "Facilitar a localização dos pontos pelas equipes;",
        "Garantir a continuidade das ações ao longo do tempo;",
        "Possibilitar o monitoramento e a avaliação da eficácia do controle.",
    ])
    add_paragraph(doc,
        "A correta identificação dos pontos é fundamental para assegurar a organização do programa, reduzir falhas operacionais e garantir maior eficiência nas aplicações de larvicida.")


def build_chapter_3(doc: Document) -> None:
    add_heading(doc, "3. Execução das Medições de Vazão em Campo", level=1, color=INK)

    add_heading(doc, "3.1 Considerações iniciais", level=2)
    add_paragraph(doc,
        "A execução das medições de vazão dos cursos d'água foi realizada conforme metodologia descrita nos itens 2.2.1 a 2.2.4 deste relatório, contemplando os pontos onde é realizada a aplicação do larvicida biológico, distribuídos nas diferentes regiões do município de Joinville.")
    add_paragraph(doc,
        "As atividades de campo foram conduzidas observando as condições hídricas e ambientais de cada ponto de aplicação, priorizando a segurança operacional e a representatividade dos dados coletados. Em locais onde foi possível executar a medição direta, todos os parâmetros hídricos (largura média, profundidade média, velocidade superficial e tempo de escoamento) foram registrados conforme a metodologia técnica adotada. Nos demais locais, em que as condições de campo inviabilizaram a medição direta, foi realizado o registro fotográfico documental e o dimensionamento da dose conduzido com base em critérios técnicos empíricos, conforme previsto na Observação Técnica do item 2.2.4.")

    add_heading(doc, "3.2 Regiões contempladas", level=2)
    add_paragraph(doc,
        "As medições e os registros foram organizados por região, conforme distribuição geográfica dos pontos de aplicação:")
    for nome, desc in [
        ("Rio do Júlio", "Medição de vazão nos cursos hídricos com pontos de aplicação distribuídos ao longo das estradas da região."),
        ("Dona Francisca", "Medição de vazão nos cursos hídricos com pontos de aplicação da Estrada Tia Marta e demais estradas da região, contemplando os cursos hídricos principais."),
        ("Região Urbana", "Medição de vazão nos pontos de aplicação localizados no Parque Zoobotânico."),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        n = p.add_run(nome + ". ")
        n.font.bold = True
        n.font.size = Pt(11)
        n.font.color.rgb = RED
        d = p.add_run(desc)
        d.font.size = Pt(11)
        d.font.color.rgb = INK

    add_heading(doc, "3.3 Pontos com medição direta de vazão", level=2)
    add_paragraph(doc,
        "Nos pontos de aplicação em que houve volume de água corrente suficiente e condições adequadas de acesso, foram executadas as medições de vazão pela metodologia descrita nos itens 2.2.1 a 2.2.3, com posterior dimensionamento da dose de larvicida conforme item 2.2.4.")
    add_paragraph(doc, "Para cada ponto avaliado, são apresentados os seguintes dados consolidados:")

    add_paragraph(doc, "Parâmetros hídricos calculados", bold=True)
    add_bullet_list(doc, [
        "Vazão (m³/s);",
        "Largura média (m);",
        "Profundidade média (m);",
        "Velocidade superficial (m/s);",
        "Dose calculada de larvicida (ml).",
    ])

    add_paragraph(doc, "Dados brutos das medições de campo", bold=True)
    add_bullet_list(doc, [
        "Larguras medidas a cada 1 metro ao longo do trecho (cm);",
        "Profundidades medidas em três pontos por seção — margem esquerda, centro e margem direita (cm);",
        "Tempos cronometrados nas repetições da medição de velocidade superficial (s).",
    ])
    add_paragraph(doc,
        "A apresentação dos dados brutos juntamente com os parâmetros calculados garante total transparência e rastreabilidade dos cálculos, permitindo a verificação técnica de cada medição apresentada.")

    add_heading(doc, "3.4 Apresentação dos pontos sem medição de vazão", level=2)
    add_paragraph(doc, "Para cada ponto sem medição direta de vazão, são apresentados nos anexos os seguintes dados:")

    add_paragraph(doc, "Dados cadastrais do ponto", bold=True)
    add_bullet_list(doc, [
        "Código de identificação do ponto;",
        "Região, estrada e cidade;",
        "Cliente vinculado;",
        "Coordenadas geográficas cadastrais;",
        "Volume atual de larvicida aplicado (ml).",
    ])

    add_paragraph(doc, "Registro fotográfico georreferenciado", bold=True)
    add_bullet_list(doc, [
        "Data e hora da captura (formato AAAA:MM:DD HH:MM:SS);",
        "Coordenadas geográficas (GPS — latitude e longitude) extraídas dos metadados EXIF;",
        "Imagem do ponto avaliado, comprovando as condições encontradas em campo.",
    ])

    add_heading(doc, "3.5 Comprovação técnica e atualizações futuras", level=2)
    add_paragraph(doc,
        "Os registros apresentados nos anexos deste relatório representam as medições e avaliações executadas até a presente data, organizadas por região e ponto de aplicação. As demais regiões do município poderão ser contempladas em medições futuras, conforme cronograma operacional e disponibilidade de condições adequadas de campo.")
    add_paragraph(doc,
        "Considerando o caráter dinâmico do programa de controle de simulídeos e a possibilidade de execução de novas medições, a relação de pontos apresentada neste relatório poderá ser atualizada e complementada em entregas futuras, mantendo-se o mesmo padrão técnico de apresentação:")
    add_bullet_list(doc, [
        "Para pontos com medição direta de vazão: parâmetros hídricos calculados, dados brutos das medições de campo e registro fotográfico georreferenciado;",
        "Para pontos sem medição direta: dados cadastrais e registro fotográfico documental georreferenciado.",
    ])
    add_paragraph(doc,
        "Esta estrutura padronizada garante a rastreabilidade e a comprovação técnica do serviço executado, tanto nas medições já realizadas quanto nas que vierem a ser incorporadas em atualizações posteriores deste produto.")


def build_signatures(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    place = doc.add_paragraph()
    place.alignment = WD_ALIGN_PARAGRAPH.CENTER
    place.paragraph_format.space_before = Pt(40)
    place.paragraph_format.space_after = Pt(60)
    p_run = place.add_run("Joinville, 29 de abril de 2026.")
    p_run.font.size = Pt(11)
    p_run.font.color.rgb = INK

    table = doc.add_table(rows=1, cols=2)
    for col_idx, (nome, cargo) in enumerate([
        ("Osmar Adelino de Aviz", "Administrador"),
        ("Eder Corbari", "Eng. Ambiental — CRQ 13302332"),
    ]):
        cell = table.rows[0].cells[col_idx]
        cell.text = ""

        line = cell.paragraphs[0]
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line.paragraph_format.space_after = Pt(2)
        line_run = line.add_run("_______________________________")
        line_run.font.size = Pt(11)
        line_run.font.color.rgb = INK

        n = cell.add_paragraph()
        n.alignment = WD_ALIGN_PARAGRAPH.CENTER
        n.paragraph_format.space_after = Pt(0)
        n_run = n.add_run(nome)
        n_run.font.bold = True
        n_run.font.size = Pt(11)
        n_run.font.color.rgb = INK

        c = cell.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c_run = c.add_run(cargo)
        c_run.font.size = Pt(10)
        c_run.font.color.rgb = MUTED


def build_annex(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "Anexos", level=1, color=INK)
    add_paragraph(doc,
        "Os anexos contêm os registros de medição por ponto e região, conforme estrutura técnica descrita nos itens 3.3 e 3.4. Material complementar disponibilizado em entregas futuras conforme cronograma.",
        italic=True, color=MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def main() -> None:
    doc = Document()
    configure_document(doc)

    build_cover(doc)
    build_toc(doc)
    build_chapter_1(doc)
    build_chapter_2(doc)
    build_chapter_3(doc)
    build_signatures(doc)
    build_annex(doc)

    out = HERE / "relatorio.docx"
    doc.save(str(out))
    print(f"OK: {out}")


if __name__ == "__main__":
    main()
